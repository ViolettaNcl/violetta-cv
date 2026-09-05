from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"
BUILD = ROOT / "build" / "cv"
SOURCE = PUBLIC / "Violetta_Nicolaou_CV.docx"
EN_DOCX = BUILD / "Violetta_Nicolaou_CV_EN.docx"
RU_DOCX = BUILD / "Violetta_Nicolaou_CV_RU.docx"
EN_PDF = PUBLIC / "Violetta_Nicolaou_CV_EN.pdf"
RU_PDF = PUBLIC / "Violetta_Nicolaou_CV_RU.pdf"


def text_nodes(paragraph):
    return paragraph._p.xpath('.//w:t')


def is_inside_hyperlink(node) -> bool:
    parent = node.getparent()
    while parent is not None and parent is not node.getroottree().getroot():
        if parent.tag == qn('w:hyperlink'):
            return True
        parent = parent.getparent()
    return False


def set_paragraph_text(paragraph, text: str) -> None:
    nodes = text_nodes(paragraph)
    if not nodes:
        paragraph.text = text
        return
    nodes[0].text = text
    for node in nodes[1:]:
        node.text = ""


def set_nonlink_prefix(paragraph, text: str) -> None:
    nodes = [node for node in text_nodes(paragraph) if not is_inside_hyperlink(node)]
    if not nodes:
        return
    nodes[0].text = text
    for node in nodes[1:]:
        node.text = ""


def set_runs(paragraph, replacements: dict[int, str]) -> None:
    for index, text in replacements.items():
        if index < len(paragraph.runs):
            paragraph.runs[index].text = text


def build_english_docx() -> None:
    doc = Document(SOURCE)
    p = doc.paragraphs

    replacements = {
        0: "VIOLETTA NICOLAOU",
        3: "Languages: Russian / English / Greek - fluent   •   French - beginner",
        4: "PROFESSIONAL PROFILE",
        5: "Junior Full-Stack .NET Developer with an honours programming diploma and a strong hands-on portfolio. I build end-to-end applications with C#, ASP.NET Core, EF Core and SQL Server, including REST APIs, JWT authentication, role-based access, real-time updates with SignalR, background services, Docker, CI/CD and automated testing. My main project is a dental-clinic web platform built for a real client. Seeking a Junior C#/.NET / ASP.NET Core / Full-Stack .NET Developer role.",
        6: "TECHNICAL STACK",
        7: "PRACTICAL DEVELOPMENT EXPERIENCE",
        9: "Designed and built a full-stack dental-clinic web platform: public website, online booking, patient portal, CRM/admin dashboard and AI assistant.",
        10: "Backend: ASP.NET Core 9, C#, REST API, layered Controllers → Services → Data architecture, EF Core 9, SQL Server, migrations and indexes.",
        11: "Implemented JWT authentication, BCrypt, role-based authorization, SignalR real-time notifications, BackgroundService, rate limiting, CORS and global error handling.",
        12: "Integrations: Google Gemini and ElevenLabs; 5-language interface; Docker, Swagger/OpenAPI, architecture/API/deployment documentation and automated checks.",
        14: "SELECTED PROJECTS",
        16: "PWA route planner with stop-order optimization (Nearest Neighbor + 2-opt), real-road OSRM routing, interactive map, weather and POI data.",
        17: "Implemented an MLP neural network with backpropagation and K-Means from scratch; added model evaluation, cross-validation, confusion matrix and precision/recall/F1.",
        18: "132 automated tests, HTTP integration and browser-flow checks, CI matrix, Docker and production smoke tests.",
        21: "Role-based desktop application for managing vehicles, drivers and routes with a structured data layer, SQL Server and automated tests.",
        22: "Produced the technical specification, architecture, class/sequence/state diagrams, developer documentation and database script.",
        25: "Multilingual responsive CV/portfolio website (EN/RU/EL) with accessible semantics, adaptive navigation and downloadable PDF CV.",
        26: "EDUCATION",
        28: "09.02.07 Information Systems and Programming - Programmer qualification, diploma with honours (red diploma).",
        29: "ADDITIONAL EXPERIENCE",
        31: "Multilingual professional communication, client service, reservations and cross-department coordination.",
    }
    for index, text in replacements.items():
        if index < len(p):
            set_paragraph_text(p[index], text)

    if len(p) > 2:
        set_nonlink_prefix(p[2], "Volgograd, Russia  •  ")
    if len(p) > 8:
        set_runs(p[8], {0: "Independent Full-Stack Developer - DentalClinic", 1: "  |  Client project · 2026"})
    if len(p) > 13:
        set_nonlink_prefix(p[13], "Demo: ")
    if len(p) > 19:
        set_nonlink_prefix(p[19], "Demo: ")
    if len(p) > 23:
        set_nonlink_prefix(p[23], "GitHub: ")
    if len(p) > 27:
        set_runs(p[27], {0: 'Moscow International College of Digital Technologies “TOP Academy”', 1: "  |  Moscow · graduated July 2026"})
    if len(p) > 30:
        set_runs(p[30], {0: "Independent Translator (2019 - Present)", 2: "Front Desk Receptionist, Crowne Plaza Limassol (2020 - 2022)"})

    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                set_paragraph_text(paragraph, "Violetta Nicolaou · Junior Full-Stack .NET Developer · github.com/ViolettaNcl")

    doc.core_properties.title = "Violetta Nicolaou — Junior Full-Stack .NET Developer"
    doc.core_properties.author = "Violetta Nicolaou"
    doc.save(EN_DOCX)


def convert_to_pdf(docx_path: Path) -> None:
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(PUBLIC), str(docx_path)],
        check=True,
    )


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Missing CV source: {SOURCE}")

    BUILD.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, RU_DOCX)
    build_english_docx()

    for output in (EN_PDF, RU_PDF):
        if output.exists():
            output.unlink()

    convert_to_pdf(RU_DOCX)
    convert_to_pdf(EN_DOCX)

    for output in (EN_PDF, RU_PDF):
        if not output.exists() or output.stat().st_size < 10_000:
            raise RuntimeError(f"PDF generation failed: {output}")
        print(f"Generated {output.relative_to(ROOT)} ({output.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
